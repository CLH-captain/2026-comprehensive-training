from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
ROOT=Path(__file__).resolve().parents[1]; ART=ROOT.parent/'artifacts'; OUT=ROOT.parent/'实训作业_Ollama与MCP工具实践.docx'
def fr(run,size=11,bold=False,color=None,name='Microsoft YaHei'):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'),name); run.font.size=Pt(size); run.bold=bold
    if color: run.font.color.rgb=RGBColor.from_string(color)
def shade(cell,fill):
    x=OxmlElement('w:shd'); x.set(qn('w:fill'),fill); cell._tc.get_or_add_tcPr().append(x)
def ct(cell,text,bold=False,color=None):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(2); fr(p.add_run(text),10,bold,color); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def code(doc,text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.right_indent=Inches(.18); p.paragraph_format.space_after=Pt(8); p.paragraph_format.line_spacing=1.05
    fr(p.add_run(text),9,name='Consolas'); x=OxmlElement('w:shd'); x.set(qn('w:fill'),'EEF5F5'); p._p.get_or_add_pPr().append(x)
def pic(doc,path,caption):
    doc.add_picture(str(path),width=Inches(6.2)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8); fr(p.add_run(caption),9,color='5A696E')
doc=Document(); sec=doc.sections[0]
for a in ('top_margin','bottom_margin','left_margin','right_margin'): setattr(sec,a,Inches(1))
sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)
normal=doc.styles['Normal']; normal.font.name='Calibri'; normal._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
for name,size,color,before,after in [('Heading 1',16,'2E74B5',16,8),('Heading 2',13,'2E74B5',12,6)]:
    st=doc.styles[name]; st.font.name='Calibri'; st._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; fr(footer.add_run('课程作业｜Ollama 本地模型与 MCP 工具实践'),9,color='6E787D')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(36); p.paragraph_format.space_after=Pt(8); fr(p.add_run('Ollama 本地大语言模型与 MCP 工具实践'),24,True,'144C56')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(24); fr(p.add_run('苏州工学院校园社团活动参与统计 Agent 课程作业'),13,color='5A696E')
p=doc.add_paragraph(); fr(p.add_run('作业目标：'),11,True); p.add_run('完成本地 Ollama 模型的安装、使用和 Python 调用，并在 Hermes Agent 中配置七个 MCP 服务，记录真实调用结果。')
doc.add_heading('一、实验环境',1)
t=doc.add_table(rows=1,cols=2); t.style='Table Grid'
for i,h in enumerate(['项目','本次使用环境']): ct(t.rows[0].cells[i],h,True,'FFFFFF'); shade(t.rows[0].cells[i],'2E74B5')
for a,b in [('操作系统','Windows 11'),('本地模型运行器','Ollama；http://127.0.0.1:11434'),('本地模型','qwen3.5-4b-64k:latest'),('Hermes','Hermes Agent CN Desktop；主模型连接本地 Ollama'),('数据库','本机 MySQL；项目数据库 szut_club_agent；MCP 只读'),('项目页面','http://127.0.0.1:5173')]:
    c=t.add_row().cells; ct(c[0],a,True); ct(c[1],b)
doc.add_heading('二、Ollama 接入本地 LLM',1); doc.add_heading('2.1 安装与模型检查',2)
doc.add_paragraph('Ollama 在本机启动大语言模型服务。通过命令行确认模型存在，再用 HTTP API 验证服务状态：')
code(doc,'D:\\ollama-windows-amd64\\ollama.exe list\n\nNAME\nqwen3.5-4b-64k:latest\nqwen3.5:4b')
doc.add_paragraph('本次使用 qwen3.5-4b-64k:latest，服务监听在 127.0.0.1:11434，演示不需要把提示词发送到云端。')
doc.add_heading('2.2 Python 调用示例',2); code(doc,(ROOT/'tools'/'ollama_python_example.py').read_text(encoding='utf-8')); doc.add_paragraph('示例使用 Python 标准库 urllib.request 调用 Ollama /api/chat 接口，以 JSON 发送模型、消息和输出参数。')
doc.add_heading('2.3 运行结果截图',2); pic(doc,ART/'ollama-result.png','图 1  Python 调用本地 Ollama 的真实返回结果')
doc.add_heading('三、MCP 工具安装与配置',1); doc.add_paragraph('MCP（Model Context Protocol）把外部工具以统一协议提供给 Hermes。Hermes 通过 stdio 启动工具进程。本次不把数据库密码写入配置文件，MySQL 连接由项目 .env 的 DATABASE_URL 动态读取。')
t=doc.add_table(rows=1,cols=3); t.style='Table Grid'
for i,h in enumerate(['MCP 服务','安装/启动方式','本项目演示用途']): ct(t.rows[0].cells[i],h,True,'FFFFFF'); shade(t.rows[0].cells[i],'2E74B5')
for row in [('Filesystem','@modelcontextprotocol/server-filesystem','读取 README 和开发文档'),('Fetch','mcp-server-fetch（Python）','抓取 Playwright 官方文档'),('Playwright','@playwright/mcp + Chromium','打开本地系统并截图'),('Git','mcp-server-git（Python）','查看 Git 状态'),('MySQL','mcp-server-mysql@1.0.42','只读查询数据库'),('Memory','@modelcontextprotocol/server-memory','保存项目知识'),('Sequential Thinking','@modelcontextprotocol/server-sequential-thinking','结构化分步思考')]:
    c=t.add_row().cells
    for i,v in enumerate(row): ct(c[i],v,i==0)
doc.add_heading('3.1 Hermes 配置截图',2); doc.add_paragraph('Hermes 配置位于本机运行目录，截图只展示服务名称和安全策略，数据库密码仍保留在项目 .env 文件中。'); pic(doc,ART/'mcp-config.png','图 2  Hermes 中 7 个 MCP 服务的本地配置记录')
doc.add_heading('3.2 MCP 使用效果截图',2); pic(doc,ART/'playwright-home.png','图 3  Playwright MCP 打开校园社团统计系统登录页'); doc.add_paragraph('Playwright MCP 实际访问 http://127.0.0.1:5173，并跟随页面跳转到登录页，证明 MCP 已能够控制本地浏览器。')
doc.add_heading('3.3 七个 MCP 的实际调用结果',2); doc.add_paragraph('逐个初始化并调用七个服务后，结果如下：'); code(doc,(ART/'mcp_demo_results.txt').read_text(encoding='utf-8',errors='replace')[-5000:]); doc.add_paragraph('验证结论：七个服务均完成初始化并返回结果。MySQL 示例使用 SELECT 查询，服务拒绝 SHOW TABLES 等 DDL 语句，说明只读限制已生效。')
doc.add_heading('四、学习与使用体会',1)
doc.add_paragraph('通过本次实践，我对本地大语言模型和 MCP 的关系有了更具体的认识。Ollama 负责在电脑上加载并运行模型，Python 可以通过 HTTP 接口发送消息并获得结构化返回，因此不依赖云端也能完成基本问答和数据分析。MCP 则把文件、网页、浏览器、Git、数据库和记忆等能力封装成统一工具，Hermes 根据模型判断选择工具，再把工具结果交回模型组织答案。实际配置中，我体会到工具的权限边界非常重要：Filesystem 只开放项目目录，Playwright 只允许访问本地页面，MySQL 使用只读查询并从本机 .env 读取密码。七个工具一起运行后，模型不只是聊天，还可以读取项目文档、查询统计表、检查代码状态并操作页面。同时，工具越多，配置和调试成本也越高，必须逐个验证启动命令、参数和返回结果。总体来看，本地模型保证隐私和可控性，MCP 提供可扩展的外部能力，两者结合适合课程项目原型，也让我理解了 Agent 系统中模型、工具和权限之间的协作关系。')
doc.add_heading('五、参考资料',1)
for item in ['Ollama API：https://github.com/ollama/ollama/blob/main/docs/api.md','MCP Servers：https://github.com/modelcontextprotocol/servers','Hermes MCP：https://hermes-agent.nousresearch.com/docs/user-guide/features/mcp/','Playwright：https://playwright.dev/docs/intro']: doc.add_paragraph(item)
for table in doc.tables:
    table.autofit=False
    for row in table.rows:
        for cell in row.cells: cell.width=Inches(6.5/len(row.cells))
OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print(OUT)

